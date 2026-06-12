import streamlit as st
import io, re, unicodedata
from datetime import datetime, date
from docx import Document

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="MobiFone Invoice – Tạo Hợp Đồng",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# DESIGN SYSTEM — MobiFone Red · Navy · White
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

* { font-family: 'Inter', sans-serif !important; box-sizing: border-box; }

/* ══ NỀN TỔNG THỂ ══ */
[data-testid="stAppViewContainer"] { background: #f0f2f5 !important; }
[data-testid="stMainBlockContainer"] { padding: 1.2rem 2rem 2rem !important; max-width: 1300px; }
[data-testid="block-container"] { padding-top: 0 !important; }

/* ══ SIDEBAR ══ */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0d2137 0%, #122d4a 100%) !important;
    border-right: none !important;
}
[data-testid="stSidebar"] section { padding: 1.2rem 1rem !important; }
[data-testid="stSidebar"] * { color: #cbd5e1 !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: #f1f5f9 !important; font-size: 0.85rem !important; font-weight: 700 !important; letter-spacing: 0.8px; text-transform: uppercase; }
[data-testid="stSidebar"] hr { border-color: #1e3a55 !important; margin: 0.8rem 0 !important; }
[data-testid="stSidebar"] label { color: #94a3b8 !important; font-size: 0.75rem !important; font-weight: 500 !important; }
[data-testid="stSidebar"] input {
    background: #1a3550 !important; color: #f1f5f9 !important;
    border: 1px solid #2a4a6b !important; border-radius: 7px !important;
    font-size: 0.83rem !important;
}
[data-testid="stSidebar"] input:focus { border-color: #e63946 !important; }
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
    background: #1a3550 !important; border: 2px dashed #2a4a6b !important;
    border-radius: 10px !important;
}

/* ══ INPUT FIELDS (main area) ══ */
[data-testid="stTextInput"] input {
    background: #ffffff !important;
    color: #0f172a !important;
    border: 1.5px solid #e2e8f0 !important;
    border-radius: 8px !important;
    font-size: 0.88rem !important;
    font-weight: 500 !important;
    padding: 0.45rem 0.75rem !important;
    transition: border-color 0.15s, box-shadow 0.15s !important;
}
[data-testid="stTextInput"] input:focus {
    border-color: #e63946 !important;
    box-shadow: 0 0 0 3px rgba(230,57,70,0.12) !important;
    outline: none !important;
}
[data-testid="stTextInput"] input::placeholder { color: #94a3b8 !important; font-weight: 400 !important; }

/* ══ LABELS ══ */
[data-testid="stTextInput"] label {
    color: #334155 !important;
    font-weight: 600 !important;
    font-size: 0.78rem !important;
    margin-bottom: 4px !important;
    letter-spacing: 0.2px !important;
}
[data-testid="stFileUploader"] label { color: #334155 !important; font-weight: 600 !important; }

/* ══ HEADER BANNER ══ */
.mbf-header {
    background: linear-gradient(135deg, #0d1f35 0%, #1a3a5c 40%, #e63946 100%);
    border-radius: 16px;
    padding: 0;
    margin-bottom: 20px;
    overflow: hidden;
    box-shadow: 0 8px 32px rgba(13,31,53,0.3);
    position: relative;
}
.mbf-header-inner {
    padding: 22px 28px;
    display: flex;
    align-items: center;
    gap: 18px;
    position: relative;
    z-index: 2;
}
.mbf-header::before {
    content: '';
    position: absolute;
    top: -40px; right: -40px;
    width: 200px; height: 200px;
    background: rgba(230,57,70,0.15);
    border-radius: 50%;
    z-index: 1;
}
.mbf-header::after {
    content: '';
    position: absolute;
    bottom: -60px; right: 80px;
    width: 160px; height: 160px;
    background: rgba(255,255,255,0.05);
    border-radius: 50%;
    z-index: 1;
}
.mbf-logo-box {
    width: 52px; height: 52px;
    background: rgba(255,255,255,0.15);
    border-radius: 14px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.8rem;
    backdrop-filter: blur(4px);
    flex-shrink: 0;
}
.mbf-header h1 {
    color: #ffffff !important;
    font-size: 1.35rem !important;
    font-weight: 800 !important;
    margin: 0 0 4px !important;
    letter-spacing: -0.3px;
}
.mbf-header .sub {
    color: rgba(255,255,255,0.65);
    font-size: 0.8rem;
    display: flex; align-items: center; gap: 10px;
}
.mbf-header .sub span {
    background: rgba(255,255,255,0.12);
    padding: 2px 10px; border-radius: 20px;
}

/* ══ STAT CARDS ══ */
.stat-row { display: flex; gap: 12px; margin-bottom: 18px; }
.stat-card {
    flex: 1;
    background: #ffffff;
    border-radius: 12px;
    padding: 14px 18px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.07);
    display: flex; align-items: center; gap: 12px;
    border: 1px solid #f1f5f9;
}
.stat-icon {
    width: 40px; height: 40px; border-radius: 10px;
    display: flex; align-items: center; justify-content: center;
    font-size: 1.2rem; flex-shrink: 0;
}
.stat-icon.red   { background: #fff1f2; }
.stat-icon.blue  { background: #eff6ff; }
.stat-icon.green { background: #f0fdf4; }
.stat-val { font-size: 1.4rem; font-weight: 800; color: #0f172a; line-height: 1; }
.stat-lbl { font-size: 0.72rem; color: #64748b; font-weight: 500; margin-top: 2px; }

/* ══ SECTION BLOCK ══ */
.sec-block {
    background: #ffffff;
    border-radius: 12px;
    margin-bottom: 14px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
    border: 1px solid #f1f5f9;
    overflow: hidden;
}
.sec-head {
    background: linear-gradient(90deg, #0d1f35, #1a3a5c);
    padding: 11px 18px;
    display: flex; align-items: center; gap: 8px;
}
.sec-head span { color: #ffffff !important; font-weight: 700; font-size: 0.82rem; letter-spacing: 0.5px; text-transform: uppercase; }
.sec-body { padding: 16px 18px 10px; }

/* ══ RIGHT PANEL CARDS ══ */
.panel-card {
    background: #ffffff;
    border-radius: 12px;
    margin-bottom: 12px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
    border: 1px solid #f1f5f9;
    overflow: hidden;
}
.panel-head {
    padding: 10px 16px;
    border-bottom: 1px solid #f1f5f9;
    font-weight: 700; font-size: 0.8rem; color: #334155;
    display: flex; align-items: center; gap: 6px;
}
.panel-body { padding: 14px 16px 12px; }

/* ══ PREVIEW FILE ══ */
.file-preview {
    background: linear-gradient(135deg, #eff6ff, #f0f9ff);
    border: 1.5px solid #bfdbfe;
    border-radius: 10px;
    padding: 12px 16px;
    display: flex; align-items: center; gap: 10px;
}
.file-preview .icon { font-size: 1.5rem; }
.file-preview .name {
    font-family: 'Courier New', monospace;
    font-size: 0.82rem; color: #1e40af;
    font-weight: 700; word-break: break-all;
}
.ma-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: #f0fdf4; border: 1px solid #bbf7d0;
    border-radius: 20px; padding: 4px 12px;
    font-size: 0.76rem; color: #166534; font-weight: 600;
    margin-top: 8px;
}

/* ══ VALIDATION ══ */
.err-list { display: flex; flex-direction: column; gap: 5px; }
.err-row {
    background: #fff5f5;
    border-left: 3px solid #e63946;
    border-radius: 0 6px 6px 0;
    padding: 6px 10px;
    font-size: 0.79rem; color: #be123c; font-weight: 500;
}
.ok-box {
    background: linear-gradient(135deg, #f0fdf4, #dcfce7);
    border: 1.5px solid #86efac;
    border-radius: 10px; padding: 12px 16px;
    text-align: center; color: #15803d;
    font-weight: 700; font-size: 0.88rem;
}

/* ══ BUTTONS ══ */
div[data-testid="stButton"] > button[kind="primary"] {
    background: linear-gradient(135deg, #c1121f, #e63946) !important;
    color: white !important; border: none !important;
    border-radius: 10px !important; font-weight: 700 !important;
    font-size: 1rem !important; letter-spacing: 0.3px !important;
    box-shadow: 0 4px 16px rgba(230,57,70,0.35) !important;
    transition: all 0.2s !important; width: 100% !important;
    padding: 0.65rem 1.5rem !important;
}
div[data-testid="stButton"] > button[kind="primary"]:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(230,57,70,0.45) !important;
}
div[data-testid="stButton"] > button[kind="primary"]:disabled {
    background: #94a3b8 !important;
    box-shadow: none !important; transform: none !important;
}

/* Secondary buttons */
div[data-testid="stButton"] > button:not([kind="primary"]) {
    background: #ffffff !important; color: #334155 !important;
    border: 1.5px solid #e2e8f0 !important; border-radius: 8px !important;
    font-weight: 600 !important; font-size: 0.83rem !important;
    transition: all 0.15s !important;
}
div[data-testid="stButton"] > button:not([kind="primary"]):hover {
    border-color: #1a3a5c !important; color: #1a3a5c !important;
    background: #f8fafc !important;
}

/* ══ DOWNLOAD BUTTON ══ */
.stDownloadButton > button {
    background: linear-gradient(135deg, #0d6832, #16a34a) !important;
    color: white !important; border: none !important;
    border-radius: 10px !important; font-weight: 700 !important;
    font-size: 1rem !important; width: 100% !important;
    box-shadow: 0 4px 16px rgba(22,163,74,0.3) !important;
    transition: all 0.2s !important;
}
.stDownloadButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(22,163,74,0.4) !important;
}

/* ══ HISTORY ITEMS ══ */
.hist-item {
    display: flex; align-items: center; gap: 8px;
    padding: 8px 10px; border-radius: 8px;
    background: #182d42; margin-bottom: 5px;
}
.hist-dot { width: 6px; height: 6px; background: #e63946; border-radius: 50%; flex-shrink: 0; }
.hist-info { flex: 1; min-width: 0; }
.hist-name { color: #e2e8f0 !important; font-size: 0.79rem; font-weight: 600; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }
.hist-time { color: #475569 !important; font-size: 0.7rem; }

/* ══ DIVIDER ══ */
hr { border: none !important; border-top: 1px solid #e2e8f0 !important; margin: 1rem 0 !important; }

/* ══ ALERTS ══ */
[data-testid="stAlert"] { border-radius: 10px !important; }

/* ══ SCROLLBAR ══ */
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #cbd5e1; border-radius: 10px; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
for k, default in [
    ("history", []),
    ("draft", {}),
    ("generated", False),
    ("result_bytes", None),
    ("result_fname", ""),
]:
    if k not in st.session_state:
        st.session_state[k] = default

# ─────────────────────────────────────────────
# CORE LOGIC
# ─────────────────────────────────────────────
def bo_dau(text: str) -> str:
    text = unicodedata.normalize("NFD", text)
    return "".join(c for c in text if unicodedata.category(c) != "Mn")

def tao_ma_hd(ten_hkd: str) -> str:
    return bo_dau(ten_hkd.upper()).replace(" ", "")

def filename_preview(so_hd: str, ten_hkd: str) -> str:
    if so_hd and ten_hkd:
        return f"HD_{so_hd}_{tao_ma_hd(ten_hkd)}.docx"
    return "HD_???.docx"

def replace_in_para(para, old: str, new: str):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    replaced = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            replaced = True
            break
    if not replaced and para.runs:
        para.runs[0].text = full.replace(old, new)
        for run in para.runs[1:]:
            run.text = ""

def get_all_paras(doc):
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras

def _xoa_ngay_ky(doc, ngay: str, thang: str, nam: str):
    ngay_str = ngay if ngay else "  "
    all_paras = get_all_paras(doc)
    for para in all_paras:
        full = "".join(r.text for r in para.runs)
        if (
            "lập vào ngày" in full
            or ("từ ngày" in full and "28" in full)
            or (
                "Ngày" in full and "tháng" in full and nam in full
                and "Luật" not in full and "Quốc Hội" not in full
            )
        ):
            for run in para.runs:
                if run.text == "28":
                    run.text = ngay_str
            replace_in_para(para, "lập vào ngày 28", f"lập vào ngày {ngay_str}")
            replace_in_para(para, "28/05/",           f"{ngay_str}/{thang}/")
            replace_in_para(para, "Ngày 28 tháng",    f"Ngày {ngay_str} tháng")

def tao_mot_hop_dong(kh: dict, template_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    email_kh = kh.get("email_kh", "")
    ep = email_kh.split("@")[0] if "@" in email_kh else email_kh
    es = ("@" + email_kh.split("@")[1]) if "@" in email_kh else ""
    ten_hkd    = kh.get("ten_hkd", "")
    ten_hkd_ng = ten_hkd[4:].strip() if ten_hkd.upper().startswith("HKD ") else ten_hkd
    ma_hd      = tao_ma_hd(ten_hkd)
    tru_so     = kh.get("tru_so", "")
    dia_chi_gd = kh.get("dia_chi_gd", "") or tru_so
    ngay       = kh.get("ngay_ky", "")
    thang      = kh.get("thang_ky", "")
    nam        = kh.get("nam_ky", "2026")

    replacements = [
        ("EMAIL_KH_PLACEHOLDER",                          ep),
        ("@GMAIL.COM",                                    es),
        ("101",                                           kh.get("so_hd", "")),
        ("HKDKIMKIEN",                                    ma_hd),
        ("NHÀ NGHỈ KIM LIÊN",                            ten_hkd_ng),
        ("142/2 Phan Châu Trinh P.Hải Châu, Tp Đà Nẵng", tru_so),
        ("0904725978",                                    kh.get("dien_thoai", "")),
        ("049189011144",                                  kh.get("ma_so_thue", "")),
        ("PHẠM THỊ KIM LIÊN",                            kh.get("dai_dien", "")),
        ("Chủ HKD",                                       kh.get("chuc_vu", "")),
        ("Trương Thị Mỹ Châu",                           kh.get("ten_nv_ben_b", "")),
        ("0901959799",                                    kh.get("sdt_ben_b", "")),
        ("cuong.danghuy.ctv@mobifone.vn",                 kh.get("email_ben_b", "")),
        ("tháng 05",                                      f"tháng {thang}"),
        ("năm 2026",                                      f"năm {nam}"),
        ("/2026",                                         f"/{nam}"),
    ]

    all_paras = get_all_paras(doc)
    for old, new in replacements:
        if old and new and old != new:
            for para in all_paras:
                replace_in_para(para, old, new)

    if dia_chi_gd != tru_so and tru_so:
        found = 0
        for para in get_all_paras(doc):
            if tru_so in "".join(r.text for r in para.runs):
                found += 1
                if found == 2:
                    replace_in_para(para, tru_so, dia_chi_gd)
                    break

    _xoa_ngay_ky(doc, ngay, thang, nam)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def validate(kh: dict) -> list:
    req_map = {
        "so_hd":        "Số hợp đồng",
        "ten_hkd":      "Tên HKD",
        "tru_so":       "Trụ sở chính",
        "dien_thoai":   "Số điện thoại",
        "ma_so_thue":   "Mã số thuế",
        "dai_dien":     "Họ tên đại diện",
        "chuc_vu":      "Chức vụ",
        "email_kh":     "Email khách hàng",
        "thang_ky":     "Tháng ký",
        "nam_ky":       "Năm ký",
        "ten_nv_ben_b": "Tên nhân viên Bên B",
        "email_ben_b":  "Email nhân viên Bên B",
        "sdt_ben_b":    "SĐT nhân viên Bên B",
    }
    errors = []
    for k, label in req_map.items():
        if not kh.get(k, "").strip():
            errors.append(f"{label}")
    if kh.get("dien_thoai") and not re.match(r"^0\d{9}$", kh["dien_thoai"]):
        errors.append("SĐT KH phải 10 số, bắt đầu bằng 0")
    if kh.get("sdt_ben_b") and not re.match(r"^0\d{9}$", kh["sdt_ben_b"]):
        errors.append("SĐT nhân viên phải 10 số, bắt đầu bằng 0")
    if kh.get("email_kh") and "@" not in kh["email_kh"]:
        errors.append("Email KH không đúng định dạng")
    if kh.get("email_ben_b") and "@" not in kh["email_ben_b"]:
        errors.append("Email nhân viên không đúng định dạng")
    return errors

# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
with st.sidebar:
    # Logo / branding sidebar
    st.markdown("""
    <div style="padding:10px 4px 14px;border-bottom:1px solid #1e3a55;margin-bottom:14px">
        <div style="font-size:1rem;font-weight:800;color:#f1f5f9;letter-spacing:0.5px">MOBIFONE AUTO-MT306</div>
        <div style="font-size:0.72rem;color:#64748b;margin-top:3px">Auto Create Contract · v7 Web</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### 📁 Template")
    template_file = st.file_uploader(
        "Upload file .docx",
        type=["docx"],
        label_visibility="collapsed",
    )

    # Upload status
    if template_file:
        st.markdown("""
        <div style="background:#0d3320;border:1px solid #166534;border-radius:8px;
                    padding:8px 12px;margin-top:6px;font-size:0.78rem;color:#4ade80;">
            ✅ Template đã sẵn sàng
        </div>""", unsafe_allow_html=True)

    st.markdown("---")

    # Link hướng dẫn sử dụng
    st.markdown("### 📖 Hướng dẫn sử dụng")
    GUIDE_URL = "https://docs.google.com/document/d/your-doc-id/edit"  # ← thay link thật vào đây
    st.markdown(f"""
    <a href="{GUIDE_URL}" target="_blank"
       style="display:flex;align-items:center;gap:8px;
              background:#1a3550;border:1px solid #2a4a6b;
              border-radius:8px;padding:9px 12px;
              text-decoration:none;margin-bottom:4px;
              transition:border-color 0.15s">
        <span style="font-size:1.1rem">📋</span>
        <div>
            <div style="color:#e2e8f0;font-size:0.8rem;font-weight:600">Xem hướng dẫn</div>
            <div style="color:#475569;font-size:0.7rem">Google Docs · Tiếng Việt</div>
        </div>
        <span style="margin-left:auto;color:#475569;font-size:0.75rem">↗</span>
    </a>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 👤 Bên B — Mặc định")
    ten_nv_ben_b = st.text_input("Tên nhân viên",  value=st.session_state.draft.get("ten_nv_ben_b", "Trương Thị Mỹ Châu"))
    email_ben_b  = st.text_input("Email nhân viên", value=st.session_state.draft.get("email_ben_b",  "cuong.danghuy.ctv@mobifone.vn"))
    sdt_ben_b    = st.text_input("SĐT nhân viên",   value=st.session_state.draft.get("sdt_ben_b",    "0901959799"))

    st.markdown("---")
    st.markdown("### 🕐 Lịch sử gần nhất")
    if not st.session_state.history:
        st.markdown("<div style='color:#334155;font-size:0.78rem;padding:4px 0'>Chưa có hợp đồng nào</div>", unsafe_allow_html=True)
    else:
        for item in reversed(st.session_state.history[-6:]):
            st.markdown(f"""
            <div class="hist-item">
                <div class="hist-dot"></div>
                <div class="hist-info">
                    <div class="hist-name">{item['fname']}</div>
                    <div class="hist-time">{item['time']}</div>
                </div>
            </div>""", unsafe_allow_html=True)
        if st.button("🗑 Xoá lịch sử", use_container_width=True):
            st.session_state.history = []
            st.rerun()

# ─────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────
st.markdown("""
<div class="mbf-header">
    <div class="mbf-header-inner">
        <div class="mbf-logo-box">📄</div>
        <div>
            <h1>MobiFone Auto Create Contract</h1>
            <div class="sub">
                <span>MOBIFONE AUTO-MT306</span>
                <span>Developed by Minh Thông · 0788 563 777</span>
                <span>v7 Web</span>
                <span>Cập nhật: 06/2026</span>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# STAT CARDS
# ─────────────────────────────────────────────
total       = len(st.session_state.history)
today_count = sum(1 for h in st.session_state.history
                  if h["time"].startswith(date.today().strftime("%d/%m/%Y")))
tpl_ok      = "✅ Đã chọn" if template_file else "❌ Chưa có"

st.markdown(f"""
<div class="stat-row">
    <div class="stat-card">
        <div class="stat-icon red">📄</div>
        <div><div class="stat-val">{total}</div><div class="stat-lbl">HĐ đã tạo</div></div>
    </div>
    <div class="stat-card">
        <div class="stat-icon blue">📅</div>
        <div><div class="stat-val">{today_count}</div><div class="stat-lbl">Hôm nay</div></div>
    </div>
    <div class="stat-card">
        <div class="stat-icon green">🗂</div>
        <div><div class="stat-val" style="font-size:0.9rem;margin-top:3px">{tpl_ok}</div><div class="stat-lbl">Template</div></div>
    </div>
</div>
""", unsafe_allow_html=True)

if not template_file:
    st.markdown("""
    <div style="background:#fff7ed;border:1.5px solid #fed7aa;border-radius:12px;
                padding:16px 20px;display:flex;align-items:center;gap:12px;">
        <span style="font-size:1.5rem">⬅️</span>
        <div>
            <div style="font-weight:700;color:#c2410c;font-size:0.9rem">Upload template để bắt đầu</div>
            <div style="color:#92400e;font-size:0.8rem;margin-top:2px">Chọn file <b>hop_dong_template.docx</b> ở sidebar bên trái</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

template_bytes = template_file.read()

# ─────────────────────────────────────────────
# FORM — 2 columns
# ─────────────────────────────────────────────
col_l, col_r = st.columns([11, 7], gap="large")

with col_l:

    # ── SECTION: THÔNG TIN HỢP ĐỒNG ──────────
    st.markdown("""
    <div class="sec-block">
        <div class="sec-head">
            <span>📋 Thông tin hợp đồng</span>
        </div>
    </div>""", unsafe_allow_html=True)

    c1, c2, c3, c4 = st.columns([2, 1.2, 1.2, 1.5])
    with c1:
        so_hd    = st.text_input("Số hợp đồng ✱",  value=st.session_state.draft.get("so_hd", ""),               placeholder="101")
    with c2:
        ngay_ky  = st.text_input("Ngày ký",         value=st.session_state.draft.get("ngay_ky",  datetime.now().strftime("%d")), placeholder="28")
    with c3:
        thang_ky = st.text_input("Tháng ký ✱",     value=st.session_state.draft.get("thang_ky", datetime.now().strftime("%m")), placeholder="06")
    with c4:
        nam_ky   = st.text_input("Năm ký ✱",       value=st.session_state.draft.get("nam_ky",   datetime.now().strftime("%Y")), placeholder="2026")

    st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

    # ── SECTION: BÊN A ───────────────────────
    st.markdown("""
    <div class="sec-block">
        <div class="sec-head">
            <span>🏪 Bên A — Khách hàng</span>
        </div>
    </div>""", unsafe_allow_html=True)

    ca1, ca2 = st.columns(2)
    with ca1:
        ten_hkd    = st.text_input("Tên HKD đầy đủ ✱",  value=st.session_state.draft.get("ten_hkd", ""),    placeholder="HKD Nhà Nghỉ Kim Liên")
    with ca2:
        tru_so     = st.text_input("Trụ sở chính ✱",     value=st.session_state.draft.get("tru_so", ""),    placeholder="142/2 Phan Châu Trinh, P.Hải Châu, Đà Nẵng")

    cb1, cb2 = st.columns(2)
    with cb1:
        dia_chi_gd = st.text_input("Địa chỉ giao dịch",  value=st.session_state.draft.get("dia_chi_gd", ""), placeholder="Để trống nếu giống trụ sở")
    with cb2:
        dien_thoai = st.text_input("Số điện thoại ✱",    value=st.session_state.draft.get("dien_thoai", ""), placeholder="0904725978")

    cc1, cc2 = st.columns(2)
    with cc1:
        ma_so_thue = st.text_input("Mã số thuế ✱",       value=st.session_state.draft.get("ma_so_thue", ""), placeholder="049189011144")
    with cc2:
        dai_dien   = st.text_input("Họ tên đại diện ✱",  value=st.session_state.draft.get("dai_dien", ""),   placeholder="PHẠM THỊ KIM LIÊN")

    cd1, cd2 = st.columns(2)
    with cd1:
        chuc_vu    = st.text_input("Chức vụ ✱",          value=st.session_state.draft.get("chuc_vu", "Chủ HKD"), placeholder="Chủ HKD")
    with cd2:
        email_kh   = st.text_input("Email khách hàng ✱", value=st.session_state.draft.get("email_kh", ""),   placeholder="kimlienpham@gmail.com")

# ── Current data dict ──
current_data = {
    "so_hd": so_hd, "ngay_ky": ngay_ky, "thang_ky": thang_ky, "nam_ky": nam_ky,
    "ten_hkd": ten_hkd, "tru_so": tru_so, "dia_chi_gd": dia_chi_gd,
    "dien_thoai": dien_thoai, "ma_so_thue": ma_so_thue,
    "dai_dien": dai_dien.upper() if dai_dien else "",
    "chuc_vu": chuc_vu, "email_kh": email_kh,
    "ten_nv_ben_b": ten_nv_ben_b, "email_ben_b": email_ben_b, "sdt_ben_b": sdt_ben_b,
}
errors      = validate(current_data)
fname_str   = filename_preview(so_hd, ten_hkd)

with col_r:

    # ── Preview file ──────────────────────────
    st.markdown('<div class="panel-card"><div class="panel-head">👁 Preview tên file</div><div class="panel-body">', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="file-preview">
        <div class="icon">📄</div>
        <div class="name">{fname_str}</div>
    </div>""", unsafe_allow_html=True)
    if ten_hkd:
        st.markdown(f'<div class="ma-badge">🔑 Mã KH: {tao_ma_hd(ten_hkd)}</div>', unsafe_allow_html=True)
    st.markdown("</div></div>", unsafe_allow_html=True)

    # ── Validation ────────────────────────────
    st.markdown('<div class="panel-card"><div class="panel-head">✅ Kiểm tra thông tin</div><div class="panel-body">', unsafe_allow_html=True)
    if errors:
        err_html = "".join(f'<div class="err-row">✕ Thiếu: {e}</div>' for e in errors)
        st.markdown(f'<div class="err-list">{err_html}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div class="ok-box">✦ Tất cả hợp lệ — Sẵn sàng tạo HĐ</div>', unsafe_allow_html=True)
    st.markdown("</div></div>", unsafe_allow_html=True)

    # ── Nháp ─────────────────────────────────
    st.markdown('<div class="panel-card"><div class="panel-head">💾 Lưu nháp</div><div class="panel-body">', unsafe_allow_html=True)
    bn1, bn2 = st.columns(2)
    with bn1:
        if st.button("💾 Lưu nháp", use_container_width=True):
            st.session_state.draft = current_data
            st.success("Đã lưu!")
    with bn2:
        if st.button("🗑 Xoá nháp", use_container_width=True):
            st.session_state.draft = {}
            st.session_state.generated = False
            st.rerun()
    st.markdown("</div></div>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# GENERATE ROW
# ─────────────────────────────────────────────
st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
gb, db = st.columns([1, 1], gap="large")

with gb:
    generate = st.button(
        "✦  Tạo Hợp Đồng",
        use_container_width=True,
        disabled=bool(errors),
        type="primary",
    )

if generate:
    try:
        result_bytes = tao_mot_hop_dong(current_data, template_bytes)
        st.session_state.result_bytes = result_bytes
        st.session_state.result_fname = fname_str
        st.session_state.generated    = True
        st.session_state.history.append({
            "fname": fname_str,
            "time":  datetime.now().strftime("%d/%m/%Y %H:%M"),
        })
        st.success(f"✅ Tạo thành công: **{fname_str}**")
    except Exception as ex:
        st.error(f"❌ Lỗi: {ex}")
        st.session_state.generated = False

if st.session_state.generated and st.session_state.result_bytes:
    with db:
        st.download_button(
            label=f"⬇️  Tải xuống {st.session_state.result_fname}",
            data=st.session_state.result_bytes,
            file_name=st.session_state.result_fname,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div style="margin-top:40px;padding:14px 0;border-top:1px solid #e2e8f0;
            display:flex;align-items:center;justify-content:space-between;">
    <div style="color:#94a3b8;font-size:0.74rem">
        © 2026 MobiFone Đà Nẵng · MOBIFONE AUTO-MT306 · Developed by Minh Thông
    </div>
    <div style="display:flex;gap:6px">
        <span style="background:#fff1f2;color:#e63946;font-size:0.7rem;
                     font-weight:700;padding:3px 10px;border-radius:20px">MobiFone Auto</span>
        <span style="background:#eff6ff;color:#1d4ed8;font-size:0.7rem;
                     font-weight:700;padding:3px 10px;border-radius:20px">MT306</span>
        <span style="background:#f0fdf4;color:#166534;font-size:0.7rem;
                     font-weight:700;padding:3px 10px;border-radius:20px">v7 Web</span>
    </div>
</div>
""", unsafe_allow_html=True)
